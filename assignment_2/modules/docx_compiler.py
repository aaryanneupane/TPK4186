import docx
from .paragraph import Paragraph
from .list import List
from .figure import Figure
from docx.shared import Inches
import os
import shutil
import tempfile

class DocxCompiler:

    def __init__(self):
        pass

    def exportDocument(self, simplanDocument, filename):
        filename_docx = filename + ".docx"
        docxDocument = docx.Document()
        self.compileDocument(simplanDocument, docxDocument)
        docxDocument.save(filename_docx)

    def compileDocument(self, simplanDocument, docxDocument):
        docxDocument.add_heading(simplanDocument.get_title(), level=1)
        self.compileSections(simplanDocument, docxDocument)

    def compileSections(self, simplanDocument, docxDocument):
        for simplanSection in simplanDocument.get_sections():
            self.compileSection(simplanSection, docxDocument)

    def compileSection(self, simplanSection, docxDocument):
        docxDocument.add_heading(simplanSection.get_title(), level=2)
        self.compileContents(simplanSection, docxDocument)

    def compileContents(self, simplanSection, docxDocument):
        for simplanContent in simplanSection.get_contents():
            if type(simplanContent) == Paragraph:
                self.compileParagraph(simplanContent, docxDocument)
            elif type(simplanContent) == List:
                if simplanContent.list_type == "uo":
                    self.compileUnorderedList(simplanContent, docxDocument)
                elif simplanContent.list_type == "ol":
                    self.compileOrderedList(simplanContent, docxDocument)
            elif type(simplanContent) == Figure:
                self.compileFigure(simplanContent, docxDocument)

    def compileParagraph(self, simplanContent, docxDocument):
        docxDocument.add_paragraph(simplanContent.get_text())

    def compileUnorderedList(self, simplanContent, docxDocument):
        for item in simplanContent.get_items():
            docxDocument.add_paragraph(item, style="ListBullet")

    def compileOrderedList(self, simplanContent, docxDocument):
        for item in simplanContent.get_items():
            docxDocument.add_paragraph(item, style="ListNumber")

    def compileFigure(self, simplanContent, docxDocument):
        image_path = simplanContent.get_path()
        if os.path.exists(image_path):
            try:
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
                    shutil.copy(image_path, temp.name)
                    docxDocument.add_picture(temp.name, width=Inches(5))
            except (IOError, SyntaxError):
                print(f"File at {image_path} is not a valid image or is corrupted.")
        else:
            print(f"No file found at {image_path}.")
